from __future__ import annotations

import base64
import email
import io
import logging
from datetime import datetime, timezone

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request as GoogleRequest
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from apps.connectors.base import BaseConnector
from apps.ingestion.document import Document

logger = logging.getLogger(__name__)

PAGE_SIZE = 50


class GmailConnector(BaseConnector):
    """
    Per-user Gmail connector.
    Each user must have their own OAuth token stored in UserSourceToken.
    """

    source = "gmail"

    def __init__(self, integration_id: str, organization_id: str) -> None:
        super().__init__(integration_id, organization_id)
        self._service = None
        self._user_email: str = ""

    def authenticate(self) -> bool:
        from apps.integrations.models import Integration
        import json

        integration = Integration.objects.get(id=self.integration_id)
        config = integration.get_config()
        token_data = config.get("token")
        if not token_data:
            return False

        creds = Credentials(
            token=token_data.get("access_token"),
            refresh_token=token_data.get("refresh_token"),
            token_uri="https://oauth2.googleapis.com/token",
            client_id=token_data.get("client_id"),
            client_secret=token_data.get("client_secret"),
            scopes=token_data.get("scopes", []),
        )

        if creds.expired and creds.refresh_token:
            creds.refresh(GoogleRequest())
            config["token"]["access_token"] = creds.token
            integration.set_config(config)
            integration.save(update_fields=["config"])

        self._service = build("gmail", "v1", credentials=creds)
        profile = self._service.users().getProfile(userId="me").execute()
        self._user_email = profile.get("emailAddress", "")
        return True

    def fetch_documents(self, cursor: dict | None = None) -> tuple[list[Document], dict]:
        if not self._service:
            self.authenticate()

        documents: list[Document] = []
        page_token = cursor.get("page_token") if cursor else None

        # Fetch messages — no label filter means all mail
        kwargs: dict = {"userId": "me", "maxResults": PAGE_SIZE}
        if page_token:
            kwargs["pageToken"] = page_token

        try:
            result = self._service.users().messages().list(**kwargs).execute()
        except HttpError as exc:
            logger.error("Gmail list error: %s", exc)
            return [], {}

        messages = result.get("messages", [])
        for msg_ref in messages:
            doc = self._fetch_message_as_document(msg_ref["id"])
            if doc:
                documents.append(doc)

        next_cursor = {}
        if result.get("nextPageToken"):
            next_cursor["page_token"] = result["nextPageToken"]

        return documents, next_cursor

    def _fetch_message_as_document(self, message_id: str) -> Document | None:
        try:
            msg = self._service.users().messages().get(
                userId="me", id=message_id, format="full"
            ).execute()
        except HttpError as exc:
            logger.warning("Could not fetch Gmail message %s: %s", message_id, exc)
            return None

        headers = {h["name"].lower(): h["value"] for h in msg.get("payload", {}).get("headers", [])}
        subject = headers.get("subject", "(no subject)")
        sender = headers.get("from", "unknown@gmail")
        date_str = headers.get("date", "")
        label_ids = msg.get("labelIds", [])

        # Parse sender
        author_email = self._parse_email_address(sender)
        author_name = sender.split("<")[0].strip().strip('"') if "<" in sender else sender

        created_at = self._parse_email_date(date_str)

        # Extract body + attachments
        body_text, attachment_names, attachment_texts = self._extract_body(msg.get("payload", {}))
        full_content = body_text
        if attachment_texts:
            full_content += "\n\n--- Attachments ---\n" + "\n\n".join(attachment_texts)

        if not full_content.strip():
            return None

        # Gmail messages are per-user (the authenticated user only sees their mail)
        from apps.accounts.models import User

        try:
            user = User.objects.get(email=self._user_email, organization_id=self.organization_id)
            allowed_user_ids = [str(user.id)]
        except User.DoesNotExist:
            allowed_user_ids = []

        return Document(
            source="gmail",
            source_item_id=message_id,
            organization_id=self.organization_id,
            title=subject,
            content=full_content,
            source_url=f"https://mail.google.com/mail/u/0/#inbox/{message_id}",
            author_email=author_email,
            author_name=author_name,
            created_at=created_at,
            updated_at=created_at,
            allowed_user_ids=allowed_user_ids,
            is_private=True,  # emails are always private
            source_metadata={
                "label_names": label_ids,
                "has_attachments": bool(attachment_names),
                "attachment_names": attachment_names,
                "thread_id": msg.get("threadId"),
                "gmail_account": self._user_email,
            },
        )

    def _extract_body(self, payload: dict) -> tuple[str, list[str], list[str]]:
        """Returns (body_text, attachment_names, attachment_texts)."""
        body_text = ""
        attachment_names = []
        attachment_texts = []

        mime_type = payload.get("mimeType", "")
        parts = payload.get("parts", [])

        if not parts:
            data = payload.get("body", {}).get("data", "")
            if data:
                decoded = base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")
                if mime_type == "text/plain":
                    body_text = decoded
                elif mime_type == "text/html":
                    body_text = self._html_to_text(decoded)

        for part in parts:
            part_mime = part.get("mimeType", "")
            filename = part.get("filename", "")

            if part_mime == "text/plain" and not filename:
                data = part.get("body", {}).get("data", "")
                if data:
                    body_text += base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")

            elif part_mime == "text/html" and not filename and not body_text:
                data = part.get("body", {}).get("data", "")
                if data:
                    html = base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")
                    body_text = self._html_to_text(html)

            elif filename:
                attachment_names.append(filename)
                # Extract text from PDF/Docx attachments
                if part_mime in ("application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                    att_text = self._extract_attachment(part)
                    if att_text:
                        attachment_texts.append(f"[{filename}]\n{att_text}")

            elif parts := part.get("parts"):
                sub_body, sub_names, sub_texts = self._extract_body(part)
                body_text = body_text or sub_body
                attachment_names.extend(sub_names)
                attachment_texts.extend(sub_texts)

        return body_text, attachment_names, attachment_texts

    def _extract_attachment(self, part: dict) -> str:
        attachment_id = part.get("body", {}).get("attachmentId")
        if not attachment_id:
            return ""
        try:
            att = self._service.users().messages().attachments().get(
                userId="me", messageId="", id=attachment_id
            ).execute()
            data = base64.urlsafe_b64decode(att["data"] + "==")
            mime_type = part.get("mimeType", "")
            if mime_type == "application/pdf":
                import pdfplumber

                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            else:
                from docx import Document as DocxDocument

                doc = DocxDocument(io.BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
        except Exception as exc:
            logger.warning("Attachment extraction failed: %s", exc)
            return ""

    def fetch_document(self, source_item_id: str) -> Document | None:
        if not self._service:
            self.authenticate()
        return self._fetch_message_as_document(source_item_id)

    def get_allowed_user_ids(self, source_item_id: str) -> list[str]:
        from apps.accounts.models import User

        try:
            user = User.objects.get(email=self._user_email, organization_id=self.organization_id)
            return [str(user.id)]
        except User.DoesNotExist:
            return []

    @staticmethod
    def _html_to_text(html: str) -> str:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n")

    @staticmethod
    def _parse_email_address(raw: str) -> str:
        if "<" in raw and ">" in raw:
            return raw.split("<")[1].rstrip(">").strip()
        return raw.strip()

    @staticmethod
    def _parse_email_date(date_str: str) -> datetime:
        try:
            from email.utils import parsedate_to_datetime

            return parsedate_to_datetime(date_str).astimezone(timezone.utc)
        except Exception:
            return datetime.now(tz=timezone.utc)
