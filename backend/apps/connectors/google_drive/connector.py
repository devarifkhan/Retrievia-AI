from __future__ import annotations

import io
import logging
from datetime import datetime, timezone

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request as GoogleRequest
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload

from apps.connectors.base import BaseConnector
from apps.ingestion.document import Document

logger = logging.getLogger(__name__)

# MIME types we can extract text from
EXPORTABLE_MIME_TYPES = {
    "application/vnd.google-apps.document": ("text/plain", ".txt"),
    "application/vnd.google-apps.spreadsheet": ("text/csv", ".csv"),
    "application/vnd.google-apps.presentation": ("text/plain", ".txt"),
}
DOWNLOADABLE_MIME_TYPES = {
    "application/pdf",
    "text/plain",
    "text/html",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

PAGE_SIZE = 100


class GoogleDriveConnector(BaseConnector):
    source = "gdrive"

    def __init__(self, integration_id: str, organization_id: str) -> None:
        super().__init__(integration_id, organization_id)
        self._service = None
        self._creds: Credentials | None = None

    def authenticate(self) -> bool:
        from apps.integrations.models import Integration
        import json

        integration = Integration.objects.get(id=self.integration_id)
        config = integration.get_config()
        token_data = config.get("token")
        if not token_data:
            logger.error("No token in Google Drive integration config")
            return False

        self._creds = Credentials(
            token=token_data.get("access_token"),
            refresh_token=token_data.get("refresh_token"),
            token_uri="https://oauth2.googleapis.com/token",
            client_id=token_data.get("client_id"),
            client_secret=token_data.get("client_secret"),
            scopes=token_data.get("scopes", []),
        )

        if self._creds.expired and self._creds.refresh_token:
            self._creds.refresh(GoogleRequest())
            # Persist refreshed token
            config["token"]["access_token"] = self._creds.token
            integration.set_config(config)
            integration.save(update_fields=["config"])

        self._service = build("drive", "v3", credentials=self._creds)
        return True

    def fetch_documents(self, cursor: dict | None = None) -> tuple[list[Document], dict]:
        if not self._service:
            self.authenticate()

        documents: list[Document] = []
        page_token = cursor.get("page_token") if cursor else None

        query = "trashed=false"
        fields = (
            "nextPageToken, files(id, name, mimeType, webViewLink, createdTime, "
            "modifiedTime, owners, parents, permissions)"
        )

        try:
            result = self._service.files().list(
                q=query,
                pageSize=PAGE_SIZE,
                fields=fields,
                pageToken=page_token,
                includeItemsFromAllDrives=True,
                supportsAllDrives=True,
            ).execute()
        except HttpError as exc:
            logger.error("Drive API error: %s", exc)
            return [], {}

        for file in result.get("files", []):
            doc = self._file_to_document(file)
            if doc:
                documents.append(doc)

        next_cursor = {}
        if result.get("nextPageToken"):
            next_cursor["page_token"] = result["nextPageToken"]

        return documents, next_cursor

    def _file_to_document(self, file: dict) -> Document | None:
        mime_type = file.get("mimeType", "")
        file_id = file["id"]

        content, is_ocr = self._extract_content(file_id, mime_type)
        if not content:
            return None

        owner = (file.get("owners") or [{}])[0]
        author_email = owner.get("emailAddress", "unknown@drive")
        author_name = owner.get("displayName", "Unknown")

        created_at = self._parse_dt(file.get("createdTime"))
        updated_at = self._parse_dt(file.get("modifiedTime"))

        # Build folder path from parents
        folder_path = self._get_folder_path(file.get("parents", []))

        # Get permissions
        allowed_user_ids = self._get_allowed_user_db_ids(file.get("permissions", []))

        return Document(
            source="gdrive",
            source_item_id=file_id,
            organization_id=self.organization_id,
            title=file.get("name", "Untitled"),
            content=content,
            source_url=file.get("webViewLink", ""),
            author_email=author_email,
            author_name=author_name,
            created_at=created_at,
            updated_at=updated_at,
            allowed_user_ids=allowed_user_ids,
            is_private=not self._is_public(file.get("permissions", [])),
            source_metadata={
                "mime_type": mime_type,
                "folder_path": folder_path,
                "is_ocr": is_ocr,
                "parents": file.get("parents", []),
            },
        )

    def _extract_content(self, file_id: str, mime_type: str) -> tuple[str, bool]:
        """Returns (text_content, is_ocr)."""
        if mime_type in EXPORTABLE_MIME_TYPES:
            export_mime, _ = EXPORTABLE_MIME_TYPES[mime_type]
            try:
                content = self._service.files().export(
                    fileId=file_id, mimeType=export_mime
                ).execute()
                return content.decode("utf-8") if isinstance(content, bytes) else content, False
            except HttpError as exc:
                logger.warning("Export failed for %s: %s", file_id, exc)
                return "", False

        if mime_type in DOWNLOADABLE_MIME_TYPES:
            data = self._download_file(file_id)
            if not data:
                return "", False

            if mime_type == "application/pdf":
                text = self._extract_pdf_text(data)
                if not text.strip():
                    text = self._ocr_pdf(data)
                    return text, True
                return text, False

            if mime_type in (
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ):
                return self._extract_docx_text(data), False

            return data.decode("utf-8", errors="replace"), False

        return "", False

    def _download_file(self, file_id: str) -> bytes | None:
        try:
            request = self._service.files().get_media(fileId=file_id)
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            return buf.getvalue()
        except HttpError as exc:
            logger.warning("Download failed for %s: %s", file_id, exc)
            return None

    def _extract_pdf_text(self, data: bytes) -> str:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)

    def _ocr_pdf(self, data: bytes) -> str:
        from .ocr import pdf_to_text_via_ocr

        return pdf_to_text_via_ocr(data)

    def _extract_docx_text(self, data: bytes) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    def fetch_document(self, source_item_id: str) -> Document | None:
        if not self._service:
            self.authenticate()
        try:
            file = self._service.files().get(
                fileId=source_item_id,
                fields="id,name,mimeType,webViewLink,createdTime,modifiedTime,owners,parents,permissions",
                supportsAllDrives=True,
            ).execute()
            return self._file_to_document(file)
        except HttpError as exc:
            if exc.resp.status == 404:
                return None
            raise

    def get_allowed_user_ids(self, source_item_id: str) -> list[str]:
        if not self._service:
            self.authenticate()
        try:
            file = self._service.files().get(
                fileId=source_item_id, fields="permissions", supportsAllDrives=True
            ).execute()
            return self._get_allowed_user_db_ids(file.get("permissions", []))
        except HttpError:
            return []

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _get_allowed_user_db_ids(self, permissions: list[dict]) -> list[str]:
        """Map Google email addresses in permissions to our DB user UUIDs."""
        from apps.accounts.models import User

        emails = {p["emailAddress"] for p in permissions if p.get("type") == "user"}
        if not emails:
            return []
        users = User.objects.filter(email__in=emails, organization_id=self.organization_id)
        return [str(u.id) for u in users]

    def _is_public(self, permissions: list[dict]) -> bool:
        return any(p.get("type") == "anyone" for p in permissions)

    def _get_folder_path(self, parent_ids: list[str]) -> str:
        if not parent_ids:
            return "/"
        try:
            folder = self._service.files().get(
                fileId=parent_ids[0], fields="name"
            ).execute()
            return f"/{folder.get('name', 'unknown')}"
        except HttpError:
            return "/"

    @staticmethod
    def _parse_dt(dt_str: str | None) -> datetime:
        if not dt_str:
            return datetime.now(tz=timezone.utc)
        return datetime.fromisoformat(dt_str.replace("Z", "+00:00"))
